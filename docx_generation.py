from docx import Document
import json
from docx.oxml import parse_xml
import tasks

def generateDocument(taskOutputFile="docx/tasks.docx", answerOutputFile="docx/answers.docx", variants=1, tasks=range(1, 19)):
    # Создание документов для заданий и ответов
    task_doc = Document()
    answer_doc = Document()
    
    # Генерация вариантов
    for variant in range(1, variants+1):
        task_doc.add_heading(f"Вариант {variant}", level=1)
        answer_doc.add_heading(f"Вариант {variant}", level=1)
        
        # Генерация задач для текущего варианта
        for i in range(len(tasks)):
            task_number = tasks[i]
            task_function = getattr(__import__(f"tasks.task{task_number}", fromlist=['']), f"generate_task{task_number}")
            task_function(task_doc, answer_doc, i+1)
        
        # Добавление разделителя между вариантами
        task_doc.add_page_break()
        answer_doc.add_page_break()
    
    # Сохранение документов
    task_doc.save(taskOutputFile)
    answer_doc.save(answerOutputFile)

for i in range(50):
    generateDocument(variants=25, tasks=[1, 2, 3, 4, 5, 6, 7, 8, 9, 10])